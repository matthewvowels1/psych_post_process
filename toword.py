import pandas as pd
from docx import Document
from docx.shared import Pt
import argparse


# takes a vtt file with speaker allocation and converts it to a docx file.

def parse_vtt(filename):
	# Initialize lists to hold the parsed data
	starts = []
	ends = []
	speakers = []
	contents = []

	# Read the file
	with open(filename, 'r') as file:
		lines = file.readlines()

	# Parse the file
	i = 0
	while i < len(lines):
		line = lines[i].strip()
		if '-->' in line:  # This is a timing line
			start, end = line.split(' --> ')
			starts.append(start.strip())
			ends.append(end.strip())

			# The next line should be the content
			i += 1
			content_line = lines[i].strip()
			if '<v ' in content_line:
				speaker, content = content_line.split('>', 1)
				speakers.append(speaker[2:].strip())
				contents.append(content.strip())
		i += 1

	# Create a DataFrame
	df = pd.DataFrame({
		'Début': starts,
		'Fin': ends,
		'Speaker': speakers,
		'Contenu': contents
	})
	return df


def create_word_document(dataframe, filename):
	# Create a new Word document
	doc = Document()
	doc.add_heading(filename, level=1)

	# Add the table
	table = doc.add_table(rows=1, cols=4)
	table.style = 'Table Grid'

	# Populate header row
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Début'
	hdr_cells[1].text = 'Fin'
	hdr_cells[2].text = 'Speaker'
	hdr_cells[3].text = 'Contenu'

	# Add the data rows
	for index, row in dataframe.iterrows():
		row_cells = table.add_row().cells
		row_cells[0].text = row['Début']
		row_cells[1].text = row['Fin']
		row_cells[2].text = row['Speaker']
		row_cells[3].text = row['Contenu']

		for cell in row_cells:
			for paragraph in cell.paragraphs:
				if not paragraph.runs:
					run = paragraph.add_run(paragraph.text)
					paragraph.text = ''
				else:
					run = paragraph.runs[0]
				run.font.size = Pt(8)

	# Save the document
	doc.save(filename)


#
def main():
	parser = argparse.ArgumentParser(description='Convert VTT file to Word Document with Table Format.')
	parser.add_argument('--input', type=str, help='Input VTT filename', required=True)
	parser.add_argument('--output', type=str, help='Output Word Document filename', required=True)
	args = parser.parse_args()

	# Parse the VTT file
	df = parse_vtt(args.input)

	# Create the Word document with the table
	create_word_document(df, args.output)


if __name__ == "__main__":
	main()
